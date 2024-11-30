import streamlit as st
from fpdf import FPDF
import pandas as pd
import json
from datetime import datetime
from pathlib import Path
import pdfkit
from docx import Document
import io

class ExportManager:
    """Gestionnaire d'exportation pour différents formats"""
    
    @staticmethod
    def to_pdf(data, title="Rapport d'Inspection"):
        """Exporte les données en PDF"""
        pdf = FPDF()
        pdf.add_page()
        
        # En-tête
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(190, 10, title, ln=True, align='C')
        pdf.ln(10)
        
        # Date
        pdf.set_font('Arial', '', 12)
        pdf.cell(190, 10, f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=True)
        
        # Contenu
        pdf.set_font('Arial', '', 12)
        if isinstance(data, dict):
            for key, value in data.items():
                pdf.cell(190, 10, f"{key}: {str(value)}", ln=True)
        elif isinstance(data, list):
            for item in data:
                pdf.cell(190, 10, str(item), ln=True)
        else:
            pdf.cell(190, 10, str(data), ln=True)
        
        # Retourne le PDF comme bytes
        return pdf.output(dest='S').encode('latin-1')
    
    @staticmethod
    def to_excel(data, sheet_name="Rapport"):
        """Exporte les données en Excel"""
        output = io.BytesIO()
        
        if isinstance(data, dict):
            df = pd.DataFrame([data])
        elif isinstance(data, list):
            df = pd.DataFrame(data)
        else:
            df = pd.DataFrame([{"data": data}])
        
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Ajustement automatique des colonnes
            worksheet = writer.sheets[sheet_name]
            for idx, col in enumerate(df.columns):
                max_length = max(df[col].astype(str).apply(len).max(),
                               len(col))
                worksheet.set_column(idx, idx, max_length + 2)
        
        return output.getvalue()
    
    @staticmethod
    def to_word(data, title="Rapport d'Inspection"):
        """Exporte les données en Word"""
        doc = Document()
        
        # Titre
        doc.add_heading(title, 0)
        
        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Contenu
        if isinstance(data, dict):
            for key, value in data.items():
                doc.add_paragraph(f"{key}: {str(value)}")
        elif isinstance(data, list):
            for item in data:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph(str(data))
        
        # Retourne le document comme bytes
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    @staticmethod
    def to_html(data, title="Rapport d'Inspection"):
        """Exporte les données en HTML"""
        html = f"""
        <html>
            <head>
                <title>{title}</title>
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        margin: 40px;
                        line-height: 1.6;
                    }}
                    h1 {{
                        color: #FF4B4B;
                        text-align: center;
                    }}
                    .date {{
                        color: #666;
                        text-align: right;
                    }}
                    .content {{
                        margin-top: 20px;
                    }}
                </style>
            </head>
            <body>
                <h1>{title}</h1>
                <div class="date">
                    {datetime.now().strftime('%d/%m/%Y %H:%M')}
                </div>
                <div class="content">
        """
        
        if isinstance(data, dict):
            for key, value in data.items():
                html += f"<p><strong>{key}:</strong> {str(value)}</p>"
        elif isinstance(data, list):
            html += "<ul>"
            for item in data:
                html += f"<li>{str(item)}</li>"
            html += "</ul>"
        else:
            html += f"<p>{str(data)}</p>"
        
        html += """
                </div>
            </body>
        </html>
        """
        
        return html
    
    @staticmethod
    def save_file(data, format="pdf", filename=None):
        """Sauvegarde les données dans un fichier"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"rapport_{timestamp}"
        
        output_dir = Path("exports")
        output_dir.mkdir(exist_ok=True)
        
        if format == "pdf":
            output_path = output_dir / f"{filename}.pdf"
            with open(output_path, "wb") as f:
                f.write(ExportManager.to_pdf(data))
        
        elif format == "excel":
            output_path = output_dir / f"{filename}.xlsx"
            with open(output_path, "wb") as f:
                f.write(ExportManager.to_excel(data))
        
        elif format == "word":
            output_path = output_dir / f"{filename}.docx"
            with open(output_path, "wb") as f:
                f.write(ExportManager.to_word(data))
        
        elif format == "html":
            output_path = output_dir / f"{filename}.html"
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(ExportManager.to_html(data))
        
        return output_path
